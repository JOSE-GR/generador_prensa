import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL


def mes_espanol(dt: datetime) -> str:
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    return meses[dt.month]


def formatear_fecha_larga() -> str:
    hoy = datetime.now()
    return f"{hoy.day} de {mes_espanol(hoy)} de {hoy.year}"


def generar_nombre_archivo() -> str:
    hoy = datetime.now()
    dia = hoy.strftime("%d")  # 06
    mes = mes_espanol(hoy)    # enero
    return f"reporte de prensa {dia} de {mes}.docx"


def generar_reporte_word(resumenes, ruta_logo) -> str:
    doc = Document()

    # --- Encabezado (tabla) ---
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    row = table.rows[0]
    row.cells[0].width = Inches(3)
    row.cells[1].width = Inches(3)

    # Reducir "aire" dentro de la tabla (sube visualmente el header)
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Logo
    if os.path.exists(ruta_logo):
        p_logo = row.cells[0].paragraphs[0]
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(0)
        p_logo.add_run().add_picture(ruta_logo, width=Inches(1.2))

    # Fecha a la derecha
    p_fecha = row.cells[1].paragraphs[0]
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fecha.paragraph_format.space_before = Pt(0)
    p_fecha.paragraph_format.space_after = Pt(0)
    run_fecha = p_fecha.add_run(formatear_fecha_larga())
    run_fecha.font.name = "Arial"
    run_fecha.font.size = Pt(17)

    # "Gerencia" centrado
    p_gg = doc.add_paragraph()
    p_gg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_gg.paragraph_format.space_before = Pt(2)
    p_gg.paragraph_format.space_after = Pt(2)
    run_gg = p_gg.add_run("Gerencia de Asuntos Económicos Internacionales")
    run_gg.font.name = "Arial"
    run_gg.font.size = Pt(17)

    # Línea divisoria (centrada)
    p_linea = doc.add_paragraph()
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_linea.paragraph_format.space_before = Pt(2)
    p_linea.paragraph_format.space_after = Pt(6)
    run_linea = p_linea.add_run("_______________________________________________")
    run_linea.font.color.rgb = RGBColor(0, 0, 0)

    # Título principal
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after = Pt(10)
    run_t = p_t.add_run("REPORTE DE PRENSA")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(28.5)

    # --- Resúmenes ---
    for item in resumenes:
        # Título de la nota
        p_titulo = doc.add_paragraph()
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after = Pt(0)

        run_titulo = p_titulo.add_run(item["titulo"])
        run_titulo.bold = True
        run_titulo.font.name = "Times New Roman"
        run_titulo.font.size = Pt(11)

        # Resumen (interlineado 1.0)
        p_resumen = doc.add_paragraph(item["resumen"])
        p_resumen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_resumen.paragraph_format.space_before = Pt(0)
        p_resumen.paragraph_format.space_after = Pt(10)

        p_resumen.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_resumen.paragraph_format.line_spacing = 1.0

        # Fuente del primer run
        if p_resumen.runs:
            run_resumen = p_resumen.runs[0]
            run_resumen.font.name = "Times New Roman"
            run_resumen.font.size = Pt(11)

    nombre_archivo = generar_nombre_archivo()
    doc.save(nombre_archivo)
    return nombre_archivo


if __name__ == "__main__":
    # Cargar resúmenes guardados
    with open("resumenes_aprobados.json", "r", encoding="utf-8") as f:
        resumenes = json.load(f)

    ruta_logo = "logo_bx.png"
    nombre = generar_reporte_word(resumenes, ruta_logo)
    print(f"✅ Documento generado: {nombre}")
