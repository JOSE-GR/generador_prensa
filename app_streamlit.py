import io
from datetime import datetime
import base64
import fitz  # PyMuPDF
import streamlit as st
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from summary_claude import resumir_con_claude


def cargar_imagen_base64(ruta: str) -> str:
    with open(ruta, "rb") as f:
        datos = f.read()
    return base64.b64encode(datos).decode("utf-8")

# ===== Marca de agua como fondo suavizado =====
watermark_b64 = cargar_imagen_base64("Logobm.svg.png")

st.markdown(
    f"""
    <style>
        .stApp {{
            /* Capa blanca semitransparente + logo de fondo */
            background-image:
                linear-gradient(rgba(255,255,255,0.96), rgba(255,255,255,0.96)),
                url("data:image/png;base64,{watermark_b64}");
            background-repeat: no-repeat;
            background-position: center 250px;  /* mueve la marca de agua hacia abajo */
            background-size: 37%;  /* ajusta 65–80 según qué tan grande lo quieras */
            background-attachment: fixed;  /* que no se mueva al hacer scroll */
        }}
    </style>
    """,
    unsafe_allow_html=True,
)



# ======== Función para agregar hipervínculos en python-docx ========

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, url, text, color="0000EE", underline=True):
    """
    Inserta un hipervínculo en un párrafo de python-docx.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )

    # Crear tag <w:hyperlink>
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Crear run <w:r>
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Color
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    # Subrayado
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    r.append(rPr)

    r_text = OxmlElement("w:t")
    r_text.text = text
    r.append(r_text)

    hyperlink.append(r)
    paragraph._p.append(hyperlink)

    return hyperlink

# ======== CONFIG MEDIOS Y FUNCIÓN PARA SEPARAR TÍTULO / MEDIO ========

MEDIOS_CONOCIDOS = {
    "Reuters",
    "Bloomberg",
    "CNBC",
    "The Wall Street Journal",
    "MarketWatch",
    "Financial Times",
    "Fox Business",
    "The Guardian",
    "PR Newswire",
}


def separar_titulo_y_medio(titulo_completo: str):
    """
    Recibe algo como:
      'ECB Officials Lobby for Rival Bank Rule Plans Before Report, Bloomberg'
    y devuelve:
      ('ECB Officials Lobby for Rival Bank Rule Plans Before Report', 'Bloomberg')

    Si no reconoce el medio, devuelve (titulo_completo, '').
    """
    t = titulo_completo.strip()

    # Quitar punto final suelto, si lo hay
    if t.endswith("."):
        t = t[:-1].strip()

    partes = [p.strip() for p in t.split(",") if p.strip()]
    if len(partes) < 2:
        return t, ""

    posible_medio = partes[-1]

    # Caso directo: la última parte es exactamente un medio conocido
    if posible_medio in MEDIOS_CONOCIDOS:
        titulo_sin_medio = ", ".join(partes[:-1]).strip()
        return titulo_sin_medio, posible_medio

    # Caso en que el medio venga mezclado, ej. 'Fox Business, Reuters'
    for medio in MEDIOS_CONOCIDOS:
        if medio in posible_medio:
            titulo_sin_medio = ", ".join(partes[:-1]).strip()
            return titulo_sin_medio, medio

    # Si no logramos identificar medio, devolvemos todo como título
    return t, ""


# ======== LÓGICA DE PDF ========
# ======== Títulos completos desde la portada (página 0) ========
def obtener_titulos_portada(doc, pagina_indice: int):
    """
    Extrae la lista de títulos de la página donde viene el índice
    (bullet points con los títulos + medio).
    """
    page = doc.load_page(pagina_indice)
    texto = page.get_text("text")

    lineas = []
    for linea in texto.splitlines():
        linea = linea.strip("• \n\t")
        if not linea:
            continue
        # Aquí puedes mantener el mismo criterio de filtrado que ya usabas
        lineas.append(linea)

    return lineas

def detectar_titulos(doc, pagina_indice: int):
    """
    1) Detecta los títulos en las páginas de los artículos (como antes),
       usando tamaño, negritas, etc. -> obtiene (titulo_interno, página).
    2) Luego busca cada titulo_interno dentro de la lista de títulos de la
       página de índice y, si lo encuentra, lo sustituye por el título
       completo que trae el periódico al final.
    """

    textos_excluidos = {"Uso General", "Información"}

    # 1) Detectar títulos internos
    titulos_raw = []

    # Las noticias empiezan justo después de la página de índice:
    #   - Si NO hay portada: índice = 0 → noticias desde 1
    #   - Si SÍ hay portada: índice = 1 → noticias desde 2
    for page_num in range(pagina_indice + 1, doc.page_count):
        page = doc.load_page(page_num)
        texto_dict = page.get_text("dict")
        blocks = texto_dict["blocks"]

        for block in blocks:
            if block["type"] != 0:
                continue

            for line in block["lines"]:
                spans = line.get("spans", [])
                if not spans:
                    continue

                full_text = "".join(span["text"] for span in spans).strip()
                font = spans[0]["font"]
                size = spans[0]["size"]
                es_negrita = "bold" in font.lower()

                if (
                    full_text
                    and es_negrita
                    and size >= 12
                    and len(full_text) >= 25
                    and full_text not in textos_excluidos
                ):
                    # Guardamos título interno y página donde inicia (1-based)
                    titulos_raw.append((full_text, page_num + 1))
                    break

    # 2) Enriquecer con los títulos completos de la página de índice
    titulos_portada = obtener_titulos_portada(doc, pagina_indice)

    titulos_enriquecidos = []
    for texto, pagina in titulos_raw:
        titulo_completo = texto
        for linea in titulos_portada:
            if texto in linea:
                titulo_completo = linea
                break
        titulos_enriquecidos.append((titulo_completo, pagina))

    return titulos_enriquecidos


def extraer_noticias_completas(doc, titulos_detectados):
    noticias = []
    for i, (titulo, pagina_inicio) in enumerate(titulos_detectados):
        if i + 1 < len(titulos_detectados):
            pagina_siguiente = titulos_detectados[i + 1][1]
            pagina_fin = pagina_siguiente - 1
        else:
            pagina_fin = doc.page_count

        texto = ""
        for num in range(pagina_inicio - 1, pagina_fin):
            page = doc.load_page(num)
            texto += page.get_text()

        noticias.append(
            {
                "titulo": titulo,
                "pagina_inicio": pagina_inicio,
                "pagina_fin": pagina_fin,
                "paginas": list(range(pagina_inicio, pagina_fin + 1)),
                "texto": texto.strip(),
            }
        )
    return noticias


def abrir_pdf_desde_bytes(pdf_bytes: bytes):
    return fitz.open(stream=pdf_bytes, filetype="pdf")


# ======== LÓGICA PARA CREAR EL WORD ========

def formatear_fecha_larga():
    hoy = datetime.now()
    dia = hoy.day
    mes_en = hoy.strftime("%B").lower()

    meses = {
        "january": "enero",
        "february": "febrero",
        "march": "marzo",
        "april": "abril",
        "may": "mayo",
        "june": "junio",
        "july": "julio",
        "august": "agosto",
        "september": "septiembre",
        "october": "octubre",
        "november": "noviembre",
        "december": "diciembre",
    }

    mes_es = meses.get(mes_en, mes_en)
    año = hoy.year
    return f"{dia} de {mes_es} de {año}"


def generar_nombre_archivo():
    hoy = datetime.now()
    dia = hoy.strftime("%d")
    mes = hoy.strftime("%B").lower().replace(" ", "")
    return f"reporte de prensa {dia}{mes}.docx"


def generar_reporte_word_en_memoria(resumenes, ruta_logo: str | None = None, pdf_url: str | None = None):
    doc = Document()

    # Márgenes un poco más amplios de texto (más ancho útil)
    for section in doc.sections:
        section.left_margin = Inches(1)   # ~2.54 cm -> puedes bajar a 0.9 si quieres aún más ancho
        section.right_margin = Inches(1)

    # Estilo por defecto del documento: Times New Roman 11
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    # Estilo de hipervínculo: si no existe, lo creamos
    try:
        hyperlink_style = doc.styles["Hyperlink"]
    except KeyError:
        hyperlink_style = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)

    hyperlink_style.font.name = "Times New Roman"
    hyperlink_style.font.size = Pt(11)

    # Encabezado con tabla
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    row = table.rows[0]
    row.cells[0].width = Inches(3)
    row.cells[1].width = Inches(3)

    # Logo
    if ruta_logo:
        try:
            row.cells[0].paragraphs[0].add_run().add_picture(
                ruta_logo, width=Inches(1.2)
            )
        except Exception:
            # si falla el logo, seguimos sin detenernos
            pass

    # Fecha
    p_fecha = row.cells[1].paragraphs[0]
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_fecha = p_fecha.add_run(formatear_fecha_larga())
    run_fecha.font.name = "Arial"
    run_fecha.font.size = Pt(17)

    # Gerencia
    p_gg = doc.add_paragraph()
    p_gg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_gg = p_gg.add_run("Gerencia de Asuntos Económicos Internacionales")
    run_gg.font.name = "Arial"
    run_gg.font.size = Pt(17)

    # Línea divisoria
    p_linea = doc.add_paragraph()
    p_linea.paragraph_format.space_before = Pt(5)
    run_linea = p_linea.add_run(
        "_____________________________________________________________"
    )
    run_linea.font.color.rgb = RGBColor(0, 0, 0)

    # Título principal
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run("REPORTE DE PRENSA")
    run_titulo.bold = True
    run_titulo.font.name = "Arial"
    run_titulo.font.size = Pt(28.5)

    doc.add_paragraph()

      # Contenido de resúmenes
    for item in resumenes:
        medio = item.get("medio", "").strip()
        titulo = item["titulo"]
        resumen = item["resumen"]
        pagina_inicio = item.get("pagina_inicio")  # puede ser None si no se llenó

        # ======= MEDIO =======
        if medio:
            p_medio = doc.add_paragraph()
            p_medio.paragraph_format.space_before = Pt(6)
            p_medio.paragraph_format.space_after = Pt(0)
            run_medio = p_medio.add_run(medio)
            run_medio.bold = True
            run_medio.font.name = "Times New Roman"
            run_medio.font.size = Pt(11)

        # ======= TÍTULO (CON HIPERVÍNCULO SI HAY URL Y PÁGINA) =======
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after = Pt(0)

        if pdf_url and pagina_inicio:
            # Construimos la URL: <pdf_url>#page=N
            link_url = f"{pdf_url}#page={pagina_inicio}"
            add_hyperlink(p_t, link_url, titulo)

            # Formato del hipervínculo (Times 11 negrita)
            for run in p_t.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                run.bold = True
        else:
            # Si no tenemos URL o página, se comporta como antes
            run_t = p_t.add_run(titulo)
            run_t.bold = True
            run_t.font.name = "Times New Roman"
            run_t.font.size = Pt(11)

        # ======= RESUMEN =======
        p_r = doc.add_paragraph(resumen)
        p_r.paragraph_format.space_before = Pt(0)
        p_r.paragraph_format.space_after = Pt(6)
        p_r.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p_r.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

        doc.add_paragraph()

    # Guardar en memoria (no en disco)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return generar_nombre_archivo(), buffer


# ======== INTERFAZ STREAMLIT ========

LOGO_PATH = "logo_bx.png"  # si está en tu carpeta

st.set_page_config(page_title="Generador de reporte de prensa", layout="wide")


# Banner superior "sticky" usando base64
banner_b64 = cargar_imagen_base64("banner_bx.png")

st.markdown(
    f"""
    <style>
    .header-banner {{
        position: sticky;
        top: 0;
        z-index: 999;
        margin-top: -50px;   /* AJUSTE: sube o baja el banner */
    }}
    </style>

    <div class="header-banner">
        <img src="data:image/png;base64,{banner_b64}"
             style="
                width: 100%;
                height: 175px;
                object-fit: cover;
                display: block;
             "/>
    </div>
    """,
    unsafe_allow_html=True,
)


st.markdown(
    """
    <h1 style="
        font-family:'Times New Roman', serif;
        font-size: 50px;
        font-weight: bold;
        text-align: center;
        margin-top: 10px;
        margin-bottom: 30px;
    ">
        Generador de Reporte de Prensa
    </h1>
    """,
    unsafe_allow_html=True
)

# === Carga del PDF y URL ===
uploaded_pdf = st.file_uploader("Sube el PDF de prensa", type=["pdf"])

pdf_url = st.text_input(
    "Pega aquí la URL del PDF (para que los títulos apunten a la página correcta):",
    value=st.session_state.get("pdf_url", ""),
)
st.session_state["pdf_url"] = pdf_url

# === Selector: ¿el PDF tiene portada? ===
tiene_portada = st.radio(
    "¿Este PDF tiene una portada (página inicial solo con portada/logo antes del índice)?",
    options=("Sí, tiene portada", "No, empieza directamente con el índice"),
    index=0,
)

if uploaded_pdf is not None:
    st.success(f"Archivo cargado: {uploaded_pdf.name}")
    pdf_bytes = uploaded_pdf.read()
    doc_pdf = abrir_pdf_desde_bytes(pdf_bytes)

    if st.button("Detectar noticias en el PDF"):
        # Si tiene portada: la página de índice es la 2 (índice=1)
        # Si no tiene portada: el índice está en la página 1 (índice=0)
        pagina_indice = 1 if tiene_portada.startswith("Sí") else 0

        titulos_detectados = detectar_titulos(doc_pdf, pagina_indice)
        if not titulos_detectados:
            st.warning("No se detectaron títulos con los criterios actuales.")
        else:
            st.session_state["titulos"] = titulos_detectados
            st.session_state["pdf_bytes"] = pdf_bytes
            st.session_state["pagina_indice"] = pagina_indice
            st.success(f"Se detectaron {len(titulos_detectados)} noticias.")


# === Selección de noticias y resúmenes ===
if "titulos" in st.session_state:
    st.subheader("Noticias detectadas")

    doc_pdf = abrir_pdf_desde_bytes(st.session_state["pdf_bytes"])
    titulos_detectados = st.session_state["titulos"]
    noticias = extraer_noticias_completas(doc_pdf, titulos_detectados)

    opciones = [f"{i+1}. {n['titulo']}" for i, n in enumerate(noticias)]
    seleccion = st.multiselect(
        "Elige las noticias que quieres resumir:",
        options=opciones,
        default=opciones,
    )

    indices = [opciones.index(op) for op in seleccion]
    noticias_seleccionadas = [noticias[i] for i in indices]

    if st.button("Generar resúmenes de las noticias seleccionadas"):
        resumenes_para_word = []
        with st.spinner("Generando resúmenes..."):
            for i, noticia in enumerate(noticias_seleccionadas, 1):
                titulo_completo = noticia["titulo"]
                titulo_nota, medio = separar_titulo_y_medio(titulo_completo)

                st.write(f"Resumiendo noticia {i}: {titulo_nota}")
                resumen = resumir_con_claude(noticia["texto"])
                st.markdown(f"**Resumen {i}:**")
                st.write(resumen)

                resumenes_para_word.append(
                    {
                        "titulo": titulo_nota,
                        "medio": medio,
                        "resumen": resumen,
                        "pagina_inicio": noticia["pagina_inicio"],
                    }
                )

        st.session_state["resumenes"] = resumenes_para_word
        st.success("Resúmenes generados.")


# === Generar y descargar el Word ===
if "resumenes" in st.session_state:
    st.subheader("Generar y descargar el reporte en Word")
    if st.button("Crear reporte de prensa"):
        nombre_archivo, buffer = generar_reporte_word_en_memoria(
            st.session_state["resumenes"],
            LOGO_PATH,
            st.session_state.get("pdf_url", ""),
        )
        st.download_button(
            label="Descargar reporte de prensa",
            data=buffer,
            file_name=nombre_archivo,
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
